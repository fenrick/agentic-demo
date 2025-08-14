"""Utilities for exporting stored lecture data to Microsoft Word (.docx)."""

from __future__ import annotations

import io
import json
import sqlite3
from typing import List

from docx import Document as DocumentFactory
from docx.document import Document

from agents.models import AssessmentItem, Citation, Slide, WeaveResult


class DocxExporter:
    """Render persisted lecture structures into a Word document."""

    def __init__(self, db_path: str) -> None:
        self._db_path = db_path

    def export(self, workspace_id: str) -> bytes:
        """Return a DOCX document for ``workspace_id`` as bytes."""

        with sqlite3.connect(self._db_path) as conn:
            lecture = self._load_lecture(conn, workspace_id)

        doc = DocumentFactory()
        self.generate_cover_page(doc, lecture)
        self.add_table_of_contents(doc)
        self.populate_sections(doc, lecture)
        self.append_references(doc, lecture.references or [])
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _load_lecture(conn: sqlite3.Connection, workspace_id: str) -> WeaveResult:
        cur = conn.execute(
            "SELECT lecture_json FROM lectures WHERE workspace_id = ? ORDER BY"
            " created_at DESC LIMIT 1",
            (workspace_id,),
        )
        row = cur.fetchone()
        cur.close()
        if row is None:
            raise ValueError("lecture not found")
        data = json.loads(row[0])
        slides = [Slide(**s) for s in data.get("slides", [])]
        assessment = [AssessmentItem(**a) for a in data.get("assessment", [])]
        references = [Citation(**c) for c in data.get("references", [])]
        return WeaveResult(
            title=data["title"],
            learning_objectives=data.get("learning_objectives", []),
            duration_min=data.get("duration_min", 0),
            author=data.get("author"),
            date=data.get("date"),
            version=data.get("version"),
            summary=data.get("summary"),
            tags=data.get("tags"),
            prerequisites=data.get("prerequisites"),
            slides=slides or None,
            assessment=assessment or None,
            references=references or None,
        )

    @staticmethod
    def generate_cover_page(doc: Document, lecture: WeaveResult) -> None:
        """Insert a simple title page based on ``lecture``."""

        doc.add_heading(lecture.title, level=0)
        if lecture.author:
            doc.add_paragraph(f"Author: {lecture.author}")
        if lecture.date:
            doc.add_paragraph(f"Date: {lecture.date}")
        if lecture.version:
            doc.add_paragraph(f"Version: {lecture.version}")
        if lecture.tags:
            doc.add_paragraph("Tags: " + ", ".join(lecture.tags))

    @staticmethod
    def add_table_of_contents(doc: Document) -> None:
        """Insert a placeholder table of contents."""

        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("TOC")

    @staticmethod
    def populate_sections(doc: Document, lecture: WeaveResult) -> None:
        """Populate ``doc`` with headings and bullet lists from ``lecture``."""

        if lecture.summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(lecture.summary)
        if lecture.learning_objectives:
            doc.add_heading("Learning Objectives", level=1)
            for obj in lecture.learning_objectives:
                doc.add_paragraph(obj, style="List Bullet")
        if lecture.prerequisites:
            doc.add_heading("Prerequisites", level=1)
            for item in lecture.prerequisites:
                doc.add_paragraph(item, style="List Bullet")
        if lecture.slides:
            for slide in lecture.slides:
                doc.add_heading(f"Slide {slide.slide_number}", level=1)
                for bullet in slide.copy.bullet_points if slide.copy else []:
                    doc.add_paragraph(bullet, style="List Bullet")
                if slide.visualization:
                    doc.add_heading("Visualisation Notes", level=2)
                    doc.add_paragraph(slide.visualization.notes)
                if slide.speaker_notes:
                    doc.add_heading("Speaker Notes", level=2)
                    doc.add_paragraph(slide.speaker_notes.notes)
        if lecture.assessment:
            doc.add_heading("Assessment", level=1)
            for item in lecture.assessment:
                desc = f"{item.type}: {item.description}"
                if item.max_score is not None:
                    desc += f" (max {item.max_score})"
                doc.add_paragraph(desc, style="List Bullet")

    @staticmethod
    def append_references(doc: Document, references: List[Citation]) -> None:
        """Append a references section for ``references``."""

        if not references:
            return
        doc.add_heading("References", level=1)
        for ref in references:
            entry = f"{ref.title} - {ref.url} (retrieved {ref.retrieved_at})"
            if ref.licence:
                entry += f" — {ref.licence}"
            doc.add_paragraph(entry, style="List Number")
