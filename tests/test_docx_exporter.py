import json
import sqlite3
from io import BytesIO
from pathlib import Path

from docx import Document

from export.docx_exporter import DocxExporter


def _setup_db(path: Path) -> Path:
    db_path = path / "test.db"
    conn = sqlite3.connect(db_path)
    conn.execute(
        "CREATE TABLE lectures (workspace_id TEXT, lecture_json TEXT, created_at TEXT)"
    )
    conn.execute(
        "CREATE TABLE citations (workspace_id TEXT, url TEXT, title TEXT, retrieved_at TEXT, licence TEXT)"
    )
    lecture = {
        "title": "Intro to AI",
        "author": "Alice",
        "date": "2024-01-01",
        "version": "1.0",
        "summary": "Basics of AI",
        "tags": ["ai", "intro"],
        "prerequisites": ["Python"],
        "duration_min": 60,
        "learning_objectives": ["Understand basics"],
        "activities": [
            {
                "type": "Lecture",
                "description": "Overview",
                "duration_min": 30,
                "learning_objectives": ["Understand basics"],
            }
        ],
        "slide_bullets": [{"slide_number": 1, "bullets": ["What is AI?", "History"]}],
        "speaker_notes": "Engage audience",
        "assessment": [{"type": "Quiz", "description": "Check", "max_score": 10}],
        "references": [
            {
                "url": "http://example.com",
                "title": "Example",
                "retrieved_at": "2024-01-01",
                "licence": "CC",
            }
        ],
    }
    conn.execute(
        "INSERT INTO lectures VALUES (?,?,?)",
        ("ws1", json.dumps(lecture), "2024-01-01"),
    )
    conn.execute(
        "INSERT INTO citations VALUES (?,?,?,?,?)",
        ("ws1", "http://example.com", "Example", "2024-01-01", "CC"),
    )
    conn.commit()
    conn.close()
    return db_path


def _doc_from_bytes(data: bytes) -> Document:
    return Document(BytesIO(data))


def test_export_returns_valid_docx_bytes(tmp_path: Path) -> None:
    db = _setup_db(tmp_path)
    exporter = DocxExporter(str(db))
    data = exporter.export("ws1")
    assert data[:2] == b"PK"
    _doc_from_bytes(data)


def test_cover_page_contents_present(tmp_path: Path) -> None:
    db = _setup_db(tmp_path)
    data = DocxExporter(str(db)).export("ws1")
    doc = _doc_from_bytes(data)
    texts = [p.text for p in doc.paragraphs]
    assert any("Author: Alice" in t for t in texts)
    assert any("Version: 1.0" in t for t in texts)
    assert any("Tags: ai, intro" in t for t in texts)


def test_sections_render_all_lecture_data(tmp_path: Path) -> None:
    db = _setup_db(tmp_path)
    data = DocxExporter(str(db)).export("ws1")
    doc = _doc_from_bytes(data)
    texts = [p.text for p in doc.paragraphs]
    assert any("Summary" in t for t in texts)
    assert any("Prerequisites" in t for t in texts)
    assert any("Slide 1" in t for t in texts)
    assert any("Speaker Notes" in t for t in texts)
    assert any("Assessment" in t for t in texts)


def test_references_section_populated(tmp_path: Path) -> None:
    db = _setup_db(tmp_path)
    data = DocxExporter(str(db)).export("ws1")
    doc = _doc_from_bytes(data)
    texts = [p.text for p in doc.paragraphs]
    assert any("References" in t for t in texts)
    assert any(
        "Example - http://example.com (retrieved 2024-01-01) — CC" in t for t in texts
    )
